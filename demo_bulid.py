from __future__ import annotations

import argparse
import json
from pathlib import Path
from typing import Iterable, List, Optional

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


DEFAULT_LOGIC_PATH = Path("chart_generation_logic.jsonl")
DEFAULT_CHART_DIR = Path("charts/generated")
DEFAULT_OUTPUT_PATH = Path("demo_report.docx")
FONT_FAMILY = "Microsoft YaHei"
EAST_ASIA_FONT_FAMILY = "微软雅黑"
TITLE_SIZE_PT = 15
CHAPTER_SIZE_PT = 14
CHART_SIZE_PT = 12
BODY_SIZE_PT = 12


def load_logic_entries(logic_path: Path) -> List[dict]:
    entries: List[dict] = []
    if not logic_path.exists():
        raise FileNotFoundError(f"Logic file not found: {logic_path}")

    with logic_path.open("r", encoding="utf-8") as file:
        for line_number, raw_line in enumerate(file, start=1):
            line = raw_line.strip()
            if not line:
                continue
            try:
                entry = json.loads(line)
            except json.JSONDecodeError as exc:
                raise ValueError(f"Invalid JSON on line {line_number}: {exc}") from exc
            entries.append(entry)
    return entries


def score_image_path(chart_id: str, path: Path) -> tuple[int, int, str]:
    """Prefer exact chart_id matches, then rank chart PNGs ahead of tables."""
    lowercase = path.name.lower()
    stem = path.stem
    match_priority = 2
    if stem == chart_id or stem.startswith(f"{chart_id}_"):
        match_priority = 0
    elif stem.startswith(chart_id):
        match_priority = 1
    chart_suffix_priority = [
        ("_pie.png", 0),
        ("_chart.png", 1),
        ("_column.png", 1),
        ("_bar.png", 1),
        ("_line.png", 1),
        ("_split.png", 1),
    ]
    chart_priority = 3
    for suffix, score in chart_suffix_priority:
        if lowercase.endswith(suffix):
            chart_priority = score
            break
    else:
        if any(keyword in lowercase for keyword in ("_pie", "_split", "_chart", "_bar", "_line", "_column")):
            chart_priority = 2
    if "table" in lowercase or "data" in lowercase:
        chart_priority = 5
    return (match_priority, chart_priority, path.name)


def find_chart_image(chart_id: str, chart_dir: Path) -> Optional[Path]:
    patterns: Iterable[str] = (
        f"{chart_id}_*.png",
        f"{chart_id}*.png",
    )
    candidates: List[Path] = []
    for pattern in patterns:
        candidates.extend(chart_dir.glob(pattern))
    if not candidates:
        return None
    return sorted(candidates, key=lambda path: score_image_path(chart_id, path))[0]


def _apply_style_font(
    document: Document, style_name: str, size_pt: Optional[int]
) -> None:
    try:
        style = document.styles[style_name]
    except KeyError:
        return

    font = style.font
    font.name = FONT_FAMILY
    if size_pt is not None:
        font.size = Pt(size_pt)

    r_pr = style._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)

    r_fonts.set(qn("w:ascii"), FONT_FAMILY)
    r_fonts.set(qn("w:hAnsi"), FONT_FAMILY)
    r_fonts.set(qn("w:cs"), FONT_FAMILY)
    r_fonts.set(qn("w:eastAsia"), EAST_ASIA_FONT_FAMILY)


def configure_document_styles(document: Document) -> None:
    """Match the template fonts for headings and body text."""
    _apply_style_font(document, "Normal", BODY_SIZE_PT)
    _apply_style_font(document, "Body Text", BODY_SIZE_PT)
    _apply_style_font(document, "Title", TITLE_SIZE_PT)
    _apply_style_font(document, "Heading 1", CHAPTER_SIZE_PT)
    _apply_style_font(document, "Heading 2", CHART_SIZE_PT)
    _apply_style_font(document, "Heading 3", BODY_SIZE_PT)


def build_document(entries: List[dict], chart_dir: Path, output_path: Path) -> None:
    document = Document()
    configure_document_styles(document)
    document.add_heading("Chart Demo Report", level=0)

    last_chapter = None
    for entry in entries:
        chapter_title = entry.get("chapter_title") or entry.get("chapter")
        if chapter_title and chapter_title != last_chapter:
            if last_chapter is not None:
                document.add_page_break()
            document.add_heading(chapter_title, level=1)
            last_chapter = chapter_title

        chart_heading = f"{entry.get('chart_id', '')} {entry.get('chart_name', '')}".strip()
        if chart_heading:
            document.add_heading(chart_heading, level=2)

        chart_id = entry.get("chart_id", "")
        image_path = find_chart_image(chart_id, chart_dir) if chart_id else None
        if image_path and image_path.exists():
            document.add_picture(str(image_path), width=Inches(6.5))
        else:
            document.add_paragraph("Image not found.")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Build a chapter-only DOCX report with charts.",
    )
    parser.add_argument(
        "--logic-path",
        type=Path,
        default=DEFAULT_LOGIC_PATH,
        help="Path to chart_generation_logic.jsonl.",
    )
    parser.add_argument(
        "--chart-dir",
        type=Path,
        default=DEFAULT_CHART_DIR,
        help="Directory containing generated chart PNG files.",
    )
    parser.add_argument(
        "--output",
        type=Path,
        default=DEFAULT_OUTPUT_PATH,
        help="Output DOCX file path.",
    )
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    entries = load_logic_entries(args.logic_path)
    build_document(entries, args.chart_dir, args.output)
    print(f"Report saved to {args.output}")


if __name__ == "__main__":
    main()

