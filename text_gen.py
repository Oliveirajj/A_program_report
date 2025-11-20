"""
主处理脚本：解析demo_report.docx，按章节生成分析文本并插入文档
"""
from __future__ import annotations

import argparse
import json
import re
from collections import defaultdict
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Union

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

from extract_docx_images import detect_heading_level
from demo_bulid import find_chart_image
from text_generator import (
    extract_template_key,
    generate_analysis_text,
    get_template_guidance,
)

# 样式配置（复用demo_bulid.py中的配置）
FONT_FAMILY = "Microsoft YaHei"
EAST_ASIA_FONT_FAMILY = "微软雅黑"
BODY_SIZE_PT = 12
DEFAULT_IMAGE_WIDTH_INCH = 6.0
DEFAULT_CHART_LOGIC_PATH = Path("chart_generation_logic.jsonl")
DEFAULT_CHART_DIR = Path("charts/generated")
MARKER_PATTERN = re.compile(
    r"[（(][^（）()]*?(?:参考图[A-Z]+|CH\d{3}[A-Z]?)(?:[^（）()]*?)[)）]"
)
MARKER_TOKEN_PATTERN = re.compile(r"(参考图[A-Z]+|CH\d{3}[A-Z]?)")


class ChapterInfo:
    """章节信息类"""
    def __init__(self, title: str, level: int, start_index: int):
        self.title = title
        self.level = level
        self.start_index = start_index
        self.content_blocks: List[Union[Paragraph, Table]] = []
        self.has_chart = False
        self.first_chart_index: Optional[int] = None


def _chapter_title_to_key(title: str, fallback: Optional[str] = None) -> str:
    """将章节标题映射为模板编号（例如3.2.1）。"""
    key = extract_template_key(title or "")
    if not key and fallback:
        key = extract_template_key(fallback) or fallback
    return key or (title or fallback or "")


def _load_chart_logic_entries(logic_path: Path) -> List[dict]:
    """读取图表逻辑定义文件"""
    entries: List[dict] = []
    if not logic_path.exists():
        print(f"警告：图表逻辑文件不存在: {logic_path}")
        return entries
    with logic_path.open("r", encoding="utf-8") as file:
        for line_no, raw_line in enumerate(file, start=1):
            line = raw_line.strip()
            if not line:
                continue
            try:
                entries.append(json.loads(line))
            except json.JSONDecodeError as exc:
                print(f"  警告：解析图表逻辑第 {line_no} 行失败: {exc}")
    return entries


def _build_chart_maps(
    logic_entries: List[dict],
    chart_dir: Path,
) -> Tuple[Dict[str, List[dict]], Dict[str, dict]]:
    """构建章节对应的图表列表，以及CH编号到图片信息的索引"""
    chapter_map: Dict[str, List[dict]] = defaultdict(list)
    chart_lookup: Dict[str, dict] = {}
    for entry in logic_entries:
        chart_id = (entry.get("chart_id") or "").strip().upper()
        if not chart_id:
            continue
        chart_name = entry.get("chart_name") or chart_id
        chapter_title = entry.get("chapter_title") or entry.get("chapter") or ""
        chapter_code = entry.get("chapter_code") or ""
        chart_info: Dict[str, str] = {
            "chart_id": chart_id,
            "chart_name": chart_name,
            "chapter_title": chapter_title,
            "chapter_code": chapter_code,
        }
        try:
            image_path = find_chart_image(chart_id, chart_dir)
        except Exception:
            image_path = None
        if image_path and image_path.exists():
            chart_info["path"] = str(image_path.resolve())
        chapter_key = _chapter_title_to_key(chapter_title, chapter_code)
        chapter_map[chapter_key].append(chart_info)
        if "path" in chart_info:
            chart_lookup[chart_id] = chart_info
    return chapter_map, chart_lookup


def _delete_paragraph(paragraph: Paragraph) -> None:
    """删除段落"""
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def remove_chapter_chart_blocks(document: Document, chapter_title: str) -> int:
    """删除章节内原有的图表标题与图片，返回删除段落数"""
    heading_index: Optional[int] = None
    for idx, para in enumerate(document.paragraphs):
        heading_info = detect_heading_level(para)
        if heading_info and heading_info[0] == 1 and heading_info[1] == chapter_title:
            heading_index = idx
            break
    if heading_index is None:
        return 0
    end_index = len(document.paragraphs)
    for idx in range(heading_index + 1, len(document.paragraphs)):
        heading_info = detect_heading_level(document.paragraphs[idx])
        if heading_info and heading_info[0] == 1:
            end_index = idx
            break
    to_remove: List[Paragraph] = []
    for para in document.paragraphs[heading_index + 1 : end_index]:
        heading_info = detect_heading_level(para)
        if heading_info and heading_info[0] >= 2 and heading_info[1].strip().upper().startswith("CH"):
            to_remove.append(para)
            continue
        if _paragraph_has_image(para):
            to_remove.append(para)
    for para in to_remove:
        _delete_paragraph(para)
    return len(to_remove)


def parse_document_chapters(doc_path: Path) -> List[ChapterInfo]:
    """
    解析文档，识别所有一级章节（Heading 1级别）
    
    Returns:
        章节信息列表
    """
    document = Document(doc_path)
    chapters: List[ChapterInfo] = []
    
    current_chapter: Optional[ChapterInfo] = None
    
    # 直接遍历document.paragraphs，这样可以直接匹配段落对象
    for para_index, para in enumerate(document.paragraphs):
        heading_info = detect_heading_level(para)
        
        if heading_info:
            level, title = heading_info
            
            # 如果是Heading 1级别，开始新章节
            if level == 1:
                # 保存上一个章节
                if current_chapter is not None:
                    chapters.append(current_chapter)
                
                # 创建新章节
                current_chapter = ChapterInfo(
                    title=title,
                    level=level,
                    start_index=para_index
                )
                # 章节标题本身也作为内容的一部分（用于后续匹配）
                current_chapter.content_blocks.append(para)
                continue
        
        # 非标题段落，添加到当前章节
        if current_chapter is not None:
            current_chapter.content_blocks.append(para)
            
            # 检查是否包含图片
            if _paragraph_has_image(para):
                if not current_chapter.has_chart:
                    current_chapter.has_chart = True
                    current_chapter.first_chart_index = len(current_chapter.content_blocks) - 1
    
    # 处理表格（表格不在paragraphs中，需要单独处理）
    # 由于表格的位置信息较难精确匹配，这里暂时跳过表格的章节归属
    # 如果需要，可以通过表格在文档中的位置来判断
    
    # 添加最后一个章节
    if current_chapter is not None:
        chapters.append(current_chapter)
    
    return chapters


def _paragraph_has_image(paragraph: Paragraph) -> bool:
    """检查段落是否包含图片"""
    for drawing in paragraph._p.iter():
        if "pic" in drawing.tag.lower() or "drawing" in drawing.tag.lower():
            return True
    return False


def extract_chapter_content(chapter: ChapterInfo, document: Document) -> str:
    """
    提取章节内容文本，用于生成分析文本
    
    Args:
        chapter: 章节信息
        document: 文档对象
        
    Returns:
        章节内容文本
    """
    content_parts = []
    
    for block in chapter.content_blocks:
        if isinstance(block, Paragraph):
            text = block.text.strip()
            # 跳过标题本身和空段落
            heading_info = detect_heading_level(block)
            if not heading_info and text:
                content_parts.append(text)
        elif isinstance(block, Table):
            # 提取表格内容
            table_text = _extract_table_text(block)
            if table_text:
                content_parts.append(f"表格：{table_text}")
    
    return "\n".join(content_parts)


def _extract_table_text(table: Table) -> str:
    """提取表格的文本内容"""
    rows_text = []
    for row in table.rows:
        cells_text = [cell.text.strip() for cell in row.cells]
        rows_text.append(" | ".join(cells_text))
    return "\n".join(rows_text)


def _prepare_analysis_paragraphs(analysis_text: str) -> List[Tuple[str, List[str]]]:
    """拆分生成文本并抽取参考图标记"""
    paragraphs: List[str] = []
    buffer: List[str] = []
    for line in analysis_text.splitlines():
        stripped = line.strip()
        if stripped:
            buffer.append(stripped)
        elif buffer:
            paragraphs.append(" ".join(buffer))
            buffer = []
    if buffer:
        paragraphs.append(" ".join(buffer))
    
    paragraph_entries: List[Tuple[str, List[str]]] = []
    for para_text in paragraphs:
        cleaned_text, markers = _extract_markers(para_text)
        paragraph_entries.append((cleaned_text, markers))
    return paragraph_entries


def _extract_markers(paragraph_text: str) -> Tuple[str, List[str]]:
    """从段落中删除(参考图X)标记并记录引用顺序"""
    markers: List[str] = []

    def _replacement(match: re.Match[str]) -> str:
        labels = MARKER_TOKEN_PATTERN.findall(match.group(0))
        markers.extend(label.strip() for label in labels)
        return ""
    
    cleaned_text = MARKER_PATTERN.sub(_replacement, paragraph_text)
    cleaned_text = re.sub(r"\s{2,}", " ", cleaned_text).strip()
    return cleaned_text, markers


def _apply_body_style(paragraph: Paragraph, document: Document) -> None:
    """统一段落字体为正文样式"""
    paragraph.style = document.styles["Normal"]
    for run in paragraph.runs:
        run.font.name = FONT_FAMILY
        run.font.size = None
        if run._element.rPr is None:
            run._element.get_or_add_rPr()
        r_fonts = run._element.rPr.rFonts
        if r_fonts is None:
            r_fonts = OxmlElement("w:rFonts")
            run._element.rPr.append(r_fonts)
        r_fonts.set(qn("w:eastAsia"), EAST_ASIA_FONT_FAMILY)


def _insert_paragraph_after(paragraph: Paragraph, text: str = "") -> Paragraph:
    """在指定段落后插入新段落"""
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    return new_para


def cleanup_orphan_markers(document: Document) -> int:
    """移除文档中遗留的(参考图X)标记并返回清理数量"""
    cleaned = 0
    for para in document.paragraphs:
        if not para.text:
            continue
        if MARKER_PATTERN.search(para.text):
            new_text = MARKER_PATTERN.sub("", para.text).strip()
            if new_text != para.text:
                para.text = new_text
                cleaned += 1
                _apply_body_style(para, document)
    return cleaned


def find_insertion_point(
    document: Document,
    chapter: ChapterInfo
) -> Tuple[Optional[Paragraph], Optional[Paragraph]]:
    """
    找到章节标题后插入文本的位置
    
    Returns:
        (插入位置的段落对象, 标题段落对象)
        如果找不到插入位置，第一个元素为None
        如果找不到标题，两个元素都为None
    """
    # 找到章节标题段落
    heading_para = None
    heading_para_index = None
    for i, para in enumerate(document.paragraphs):
        heading_info = detect_heading_level(para)
        if heading_info and heading_info[1] == chapter.title:
            heading_para = para
            heading_para_index = i
            break
    
    if heading_para is None:
        return None, None
    
    # 查找标题后的第一个空段落或第一个图表/子标题前的位置
    # 如果章节有图表，在第一个图表或其子标题前插入
    # 否则在标题后的第一个空段落插入
    
    text_seen = False
    # 先检查标题后的段落
    for i in range(heading_para_index + 1, len(document.paragraphs)):
        para = document.paragraphs[i]
        
        # 如果遇到下一个一级标题，停止
        heading_info = detect_heading_level(para)
        if heading_info and heading_info[0] == 1:
            break

        # 如果遇到二级或三级标题且此前没有正文，则在其前插入文本
        if heading_info and heading_info[0] >= 2 and not text_seen:
            return para, heading_para

        # 如果遇到空段落，可以在这里插入（在它之前插入）
        if not para.text.strip():
            return para, heading_para
        
        # 如果这个段落包含图片，在它之前插入
        if _paragraph_has_image(para):
            return para, heading_para

        # 标记已遇到正文内容
        if not heading_info and para.text.strip():
            text_seen = True
    
    # 如果没找到合适位置，在标题后的第一个段落前插入
    # 如果标题后没有段落，则在标题后添加新段落
    if heading_para_index + 1 < len(document.paragraphs):
        return document.paragraphs[heading_para_index + 1], heading_para
    else:
        # 标题是最后一个段落，返回None，将在标题后添加
        return None, heading_para


def insert_analysis_text(
    document: Document,
    chapter: ChapterInfo,
    analysis_text: str,
    chart_lookup: Dict[str, dict],
) -> None:
    """
    在文档中插入分析文本
    
    Args:
        document: 文档对象
        chapter: 章节信息
        analysis_text: 要插入的分析文本
    """
    paragraphs_with_markers = _prepare_analysis_paragraphs(analysis_text)
    if not paragraphs_with_markers:
        print(f"  提示：章节 '{chapter.title}' 生成文本为空，跳过插入")
        return

    insert_para, heading_para = find_insertion_point(document, chapter)
    
    if heading_para is None:
        print(f"警告：无法找到章节 '{chapter.title}' 的标题，跳过")
        return
    
    inserted_entries: List[Tuple[Paragraph, List[str]]] = []
    if insert_para is None:
        anchor = heading_para
        for text, markers in paragraphs_with_markers:
            new_para = _insert_paragraph_after(anchor, text)
            anchor = new_para
            inserted_entries.append((new_para, markers))
    else:
        anchor = insert_para
        temp_entries: List[Tuple[Paragraph, List[str]]] = []
        for text, markers in reversed(paragraphs_with_markers):
            new_para = anchor.insert_paragraph_before(text)
            temp_entries.append((new_para, markers))
        inserted_entries = list(reversed(temp_entries))
    
    for paragraph, markers in inserted_entries:
        _apply_body_style(paragraph, document)
        if not markers:
            continue
        current_anchor = paragraph
        for marker in markers:
            image_info = chart_lookup.get(marker.upper())
            if not image_info:
                print(f"  警告：未找到{marker}对应的CH图表，跳过插入")
                continue
            image_path = image_info.get("path")
            if not image_path:
                print(f"  警告：{marker}缺少可用的图片路径，跳过插入")
                continue
            image_para = _insert_paragraph_after(current_anchor)
            _apply_body_style(image_para, document)
            try:
                run = image_para.add_run()
                run.add_picture(image_path, width=Inches(DEFAULT_IMAGE_WIDTH_INCH))
                image_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as exc:
                label = image_info.get("name") or image_info.get("chart_name") or "未知图片"
                print(f"  警告：插入{marker}（{label}）失败: {exc}")
                # 插入失败时移除空段落
                parent = image_para._element.getparent()
                if parent is not None:
                    parent.remove(image_para._element)
                continue
            current_anchor = image_para


def process_document(
    input_path: Path,
    output_path: Path,
    chart_logic_path: Path = DEFAULT_CHART_LOGIC_PATH,
    chart_dir: Path = DEFAULT_CHART_DIR,
) -> None:
    """
    处理文档：解析章节，生成分析文本，插入文档
    
    Args:
        input_path: 输入文档路径（demo_report.docx）
        output_path: 输出文档路径
        chart_logic_path: 图表逻辑定义文件路径
        chart_dir: 图表图片目录
    """
    print(f"正在加载文档: {input_path}")
    document = Document(input_path)
    
    print("正在解析章节...")
    chapters = parse_document_chapters(input_path)
    print(f"找到 {len(chapters)} 个一级章节")
    
    # 加载章节图表映射
    print(f"正在读取图表逻辑: {chart_logic_path}")
    chart_logic_entries = _load_chart_logic_entries(chart_logic_path)
    chapter_chart_map, chart_lookup = _build_chart_maps(chart_logic_entries, chart_dir)
    print(f"已关联 {len(chart_lookup)} 个CH图表资源")
    
    # 为每个章节生成分析文本
    for i, chapter in enumerate(chapters, 1):
        print(f"\n处理章节 {i}/{len(chapters)}: {chapter.title}")
        
        # 移除模板中遗留的图表及标题
        removed = remove_chapter_chart_blocks(document, chapter.title)
        if removed:
            print(f"  已移除 {removed} 个模板图表段落")

        # 提取章节内容
        chapter_content = extract_chapter_content(chapter, document)
        
        if not chapter_content.strip():
            template_guidance = get_template_guidance(chapter.title)
            if template_guidance:
                chapter_content = template_guidance["text"]
                print("  章节内容为空，使用模板参考文本作为生成依据")
            else:
                print(f"  章节内容为空且无模板参考，跳过")
                continue
        
        # 生成分析文本
        print(f"  正在生成分析文本...")
        chapter_key = _chapter_title_to_key(chapter.title)
        chapter_charts = chapter_chart_map.get(chapter_key, [])
        analysis_text = generate_analysis_text(
            chapter_title=chapter.title,
            chapter_content=chapter_content,
            reference_images=[],
            chapter_charts=chapter_charts,
        )
        
        if analysis_text and not analysis_text.startswith("[文本生成失败"):
            print(f"  生成成功，长度: {len(analysis_text)} 字符")
            # 插入文档
            insert_analysis_text(
                document,
                chapter,
                analysis_text,
                chart_lookup,
            )
        else:
            print(f"  生成失败，跳过")
    
    cleaned_markers = cleanup_orphan_markers(document)
    if cleaned_markers:
        print(f"\n清理了 {cleaned_markers} 个残留的参考图标记")
    
    # 保存文档
    print(f"\n正在保存文档: {output_path}")
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))
    print(f"完成！文档已保存到: {output_path}")


def main():
    """主函数"""
    parser = argparse.ArgumentParser(
        description="为demo_report.docx的每个章节生成分析文本",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例：
  # 使用默认路径
  python generate_analysis_text.py
  
  # 指定输入和输出文件
  python generate_analysis_text.py -i demo_report.docx -o 能耗报告_自动生成.docx
  
  # 指定参考图片文件夹
  python generate_analysis_text.py --ref-dir ref
        """
    )
    
    parser.add_argument(
        "-i", "--input",
        type=Path,
        default=Path("demo_report.docx"),
        help="输入的DOCX文件路径（默认: demo_report.docx）"
    )
    
    parser.add_argument(
        "-o", "--output",
        type=Path,
        default=Path("能耗报告_自动生成.docx"),
        help="输出的DOCX文件路径（默认: 能耗报告_自动生成.docx）"
    )
    
    parser.add_argument(
        "--chart-logic",
        type=Path,
        default=DEFAULT_CHART_LOGIC_PATH,
        help="图表逻辑定义文件路径（默认: chart_generation_logic.jsonl）",
    )

    parser.add_argument(
        "--chart-dir",
        type=Path,
        default=DEFAULT_CHART_DIR,
        help="CH图表图片所在目录（默认: charts/generated）",
    )
    
    args = parser.parse_args()
    
    # 检查输入文件
    if not args.input.exists():
        print(f"错误：输入文件不存在: {args.input}")
        return 1
    
    try:
        process_document(
            args.input,
            args.output,
            chart_logic_path=args.chart_logic,
            chart_dir=args.chart_dir,
        )
        return 0
    except Exception as e:
        print(f"错误：处理文档时发生异常: {e}")
        import traceback
        traceback.print_exc()
        return 1


if __name__ == "__main__":
    exit(main())

